from io import BytesIO
import zlib
from bson import ObjectId
from transformers import AutoProcessor, AutoModelForVision2Seq #, TrOCRProcessor, VisionEncoderDecoderModel
import os
import re
import PyPDF2
import pdfplumber
from pptxtopdf import convert as pptx_to_pdf
from docx2pdf import convert as docx_to_pdf
from app.models.user import TextChunk, Image
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image as PilImage
from app.config import Config


class OCRService:
    def __init__(self):
        # Nougat Modell und Prozessor laden
        self.processor = AutoProcessor.from_pretrained("facebook/nougat-small")
        self.model = AutoModelForVision2Seq.from_pretrained("facebook/nougat-small")
        # self.processorFormula = TrOCRProcessor.from_pretrained("microsoft/trocr-small-handwritten")
        # self.modelFormula = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-small-handwritten")

    def extract_text(self, doc):
        file_extension = os.path.splitext(doc.name)[1].lower()

        if file_extension == ".pdf":
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        elif file_extension in [".docx", ".pptx"]:
            self.convert_docx_or_pptx_to_pdf(doc)  # Konvertiere DOCX/PPTX zu PDF
            self.process_pdf(doc)  # PDF-Verarbeitung aufrufen
        else:
            print("Unsupported file type")  # Überprüfen, ob der Dateityp nicht unterstützt wird

        return doc


# Convert
    def convert_docx_or_pptx_to_pdf(self, doc):
        try:
            if doc.name.endswith(".docx"):
                docx_to_pdf(doc.doclink)  # Convert DOCX to PDF
            elif doc.name.endswith(".pptx"):
                pptx_to_pdf(doc.doclink)  # Convert PPTX to PDF
        except Exception as e:
            raise Exception(f"Error during conversion: {e}")

        doc.name = doc.name.replace('.docx', '.pdf').replace('.pptx', '.pdf')
        return doc

    def convert_docx_to_pdf(self, doc):
        print(f"Starting DOCX to PDF conversion for: {doc.doclink}")
        
        try:
            doc_content = DocxDocument(doc.doclink)
            print(f"Number of paragraphs in DOCX: {len(doc_content.paragraphs)}")
            
            pdf_stream = BytesIO()
            pdf_canvas = canvas.Canvas(pdf_stream, pagesize=letter)

            for idx, paragraph in enumerate(doc_content.paragraphs):
                print(f"Writing Paragraph {idx + 1} to PDF: {paragraph.text}")
                pdf_canvas.drawString(100, 750, paragraph.text)
                pdf_canvas.showPage()

            pdf_canvas.save()
            pdf_stream.seek(0)

            pdf_length = pdf_stream.getbuffer().nbytes
            print(f"PDF Stream created with length: {pdf_length} bytes")

            doc.doclink = pdf_stream
            return pdf_stream

        except Exception as e:
            print(f"Error during DOCX to PDF conversion: {e}")
            raise


    def convert_pptx_to_pdf(self, doc):
        # Dateiinhalt in Bytes laden
        pptx_stream = BytesIO(doc.doclink.read())
        pdf_stream = BytesIO()  # Erstelle einen neuen BytesIO-Stream für die PDF

        try:
            # Konvertiere PPTX-Stream direkt zu PDF-Stream
            pptx_to_pdf(pptx_stream, pdf_stream)  # Anpassung hier
            pdf_stream.seek(0)  # Zurücksetzen des Streams auf den Anfang
        except Exception as e:
            raise Exception(f"Fehler bei der Konvertierung von PPTX zu PDF: {e}")

        return pdf_stream  # Gibt den PDF-Stream zurück

# handle formulas
    # def extract_formulas_from_image(self, image_path):
    #     processorFormula = TrOCRProcessor.from_pretrained("microsoft/trocr-small-handwritten")
    #     modelFormula = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-small-handwritten")

    #     image = Image.open(image_path).convert("RGB")
    #     pixel_values = processorFormula(image, return_tensors="pt").pixel_values
    #     generated_ids = modelFormula.generate(pixel_values)
    #     generated_text = processorFormula.batch_decode(generated_ids, skip_special_tokens=True)[0]
        
    #     # Example regex to identify LaTeX-like formulas from the recognized text
    #     formulas = re.findall(r'\$[^\$]+\$', generated_text)
    #     return formulas
    
    # def save_formulas_to_file(self, formulas, output_path):
    #     with open(output_path, "w", encoding="utf-8") as f:
    #         for formula, page_num in formulas:
    #             f.write(f"Formula on Page {page_num}: {formula}\n")

    # def extract_text_formulas_from_pdf(self, file_path):
    #     formulas = []
    #     with pdfplumber.open(file_path) as pdf:
    #         for page_num, page in enumerate(pdf.pages):
    #             text = page.extract_text()
    #             matches = re.findall(r'\$[^\$]+\$', text)  # Extract LaTeX-like formulas
    #             formulas.extend([(match, page_num + 1) for match in matches])
    #     return formulas


# extraction
    def extract_text_from_pdf(self, pdf_stream):
        full_text = ""
        page_texts = []  # Liste, um Text pro Seite zu speichern
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    full_text += f"{page_text.strip()}\n"
                    page_texts.append((page_text.strip(), page_num + 1))  # Seitenzahl hinzufügen
                else:
                    print(f"No text found on page {page_num + 1}.")
                    full_text += f"No text detected on page {page_num + 1}\n"
                    page_texts.append(("", page_num + 1))  # Leeren Text hinzufügen
            
            if not full_text.strip():  # If no text was found
                print("The document appears to contain only images.")
                full_text = ""  # Return empty string for image-only documents

        except Exception as e:
            print(f"Error extracting text from PDF: {e}")

        return page_texts  # Gib die Liste von Seiten zurück


    def extract_images_from_pdf(self, pdf_stream, doc):
        images = []
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            for page_num, page in enumerate(pdf_reader.pages):
                resources = page.get("/Resources").get_object()
                xobjects = resources.get("/XObject")
                
                if xobjects:
                    xobjects = xobjects.get_object()
                    image_counter = 1  # Image counter für jede Seite zurücksetzen
                    
                    for obj in xobjects:
                        xobject = xobjects[obj].get_object()
                        
                        if xobject["/Subtype"] == "/Image":
                            try:
                                width = xobject["/Width"]
                                height = xobject["/Height"]

                                # Daten extrahieren
                                data = xobject._data
                                file_extension = "png"  # Standard PNG verwenden

                                # Überprüfe auf vorhandene Filter und dekodiere entsprechend
                                if "/Filter" in xobject:
                                    if xobject["/Filter"] == "/FlateDecode":
                                        try:
                                            data = zlib.decompress(data)  # Dekomprimiere die FlateDecode-Daten
                                            # Erstelle ein Bild aus den dekomprimierten Daten
                                            img = PilImage.frombytes("RGB", (width, height), data)
                                        except Exception as e:
                                            continue
                                    elif xobject["/Filter"] == "/DCTDecode":
                                        file_extension = "jpg"  # JPEG benötigt keine Decodierung
                                    elif xobject["/Filter"] == "/JPXDecode":
                                        file_extension = "jp2"  # JPEG2000
                                    else:
                                        continue  # Überspringe unbekannte Filter

                                # Speicherpfad für das Bild
                                img_save_path = os.path.join(Config.LOCAL_DOC_PATH, f"extracted_image_{page_num+1}_{image_counter}.{file_extension}")

                                # Speichere das Bild basierend auf dem Filtertyp
                                img.save(img_save_path)

                                # OCR auf dem Bild ausführen
                                print(f"Starte OCR für Bild {image_counter} auf Seite {page_num + 1}...")
                                img_text = self.ocr_image(img_save_path)  # OCR-Funktion

                                # Erstelle ein Image-Objekt
                                image_obj = Image(
                                    id=str(image_counter),  # oder eine andere Logik für die ID
                                    link=img_save_path,
                                    page=page_num + 1,
                                    type=file_extension,
                                    imgtext=img_text,  # Hier die von OCR zurückgegebene Text
                                    llm_output=""  # Platzhalter für LLM-Ausgabe
                                )
                                images.append(image_obj)

                                image_counter += 1

                            except Exception as e:
                                print(f"Fehler beim Verarbeiten des Bildes für Objekt {obj} auf Seite {page_num + 1}: {e}")

        except Exception as e:
            print(f"Fehler beim Extrahieren der Bilder aus dem PDF: {e}")

        print("Bildextraktion abgeschlossen.")
        return images


    def split_text_into_chunks(self, full_text, doc, page_num):
            # Split based on double newlines to capture paragraphs or sections
            chunks = full_text.split('\n\n')
            for chunk in chunks:
                if chunk.strip():
                    text_chunk = TextChunk(id=str(ObjectId()), text=chunk.strip(), page=page_num)
                    doc.textList.append(text_chunk)

    def ocr_image(self, image_stream):
        try:            
            # Lade das Bild aus dem Stream
            image = PilImage.open(image_stream).convert("RGB")
            
            # Vorverarbeitung des Bildes für das Nougat-Modell
            pixel_values = self.processor(images=image, return_tensors="pt").pixel_values
            
            # OCR durchführen mit dem Nougat-Modell
            generated_ids = self.model.generate(pixel_values)
            generated_text = self.processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
            
            return generated_text  # Gib nur den erkannten Text zurück
            
        except Exception as e:
            print(f"Fehler bei der OCR für Bild: {e}")
            return None

# Sudo main
    def process_pdf(self, doc):
        print(f"Starting PDF processing for document: {doc.name}")

        try:
            # Attempt to extract text from PDF
            print("Attempting to extract text from PDF...")
            page_texts = self.extract_text_from_pdf(doc.doclink)
            print(f"Extracted text from PDF.")

            # Split extracted text into chunks
            print("Splitting extracted text into chunks...")
            for text, page_num in page_texts:
                self.split_text_into_chunks(text, doc, page_num)  # page_num wird jetzt korrekt übergeben
            print(f"Text chunks created: {len(doc.textList)}")

            # Always attempt image extraction regardless of text outcome
            print("Attempting to extract images from PDF...")
            self.extract_images_from_pdf(doc.doclink, doc)
            print(f"Image extraction complete. Images found: {len(doc.imgList)}")

        except Exception as e:
            print(f"Error during PDF processing: {e}")

